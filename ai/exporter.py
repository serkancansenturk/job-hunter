"""
TailoredCV'yi DOCX ve PDF olarak dışa aktarır.
"""
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config.cv_base import CV_DATA
from models.cv import TailoredCV


OUTPUT_DIR = Path(__file__).parent.parent / "output"


class CVExporter:
    def __init__(self):
        OUTPUT_DIR.mkdir(exist_ok=True)

    def export_docx(self, cv: TailoredCV) -> str:
        doc = Document()
        self._set_margins(doc)
        self._add_header(doc, cv)
        self._add_summary(doc, cv)
        self._add_experience(doc, cv)
        self._add_skills(doc)
        self._add_certifications(doc)
        self._add_education(doc)

        slug = f"{cv.company.replace(' ', '_')}_{cv.job_title.replace(' ', '_')[:30]}"
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        filename = OUTPUT_DIR / f"CV_Serkan_{slug}_{ts}.docx"
        doc.save(filename)
        cv.docx_path = str(filename)
        return str(filename)

    def _set_margins(self, doc: Document) -> None:
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

    def _add_header(self, doc: Document, cv: TailoredCV) -> None:
        personal = CV_DATA["personal"]
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = name_p.add_run(personal["name"].upper())
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1a, 0x20, 0x2c)

        title_p = doc.add_paragraph()
        run2 = title_p.add_run(cv.job_title)
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

        linkedin = personal.get("linkedin", "")
        contact_line = f"{personal['phone']}  ·  {personal['email']}  ·  {personal['location']}"
        if linkedin:
            contact_line += f"  ·  {linkedin}"

        contact_p = doc.add_paragraph()
        contact_p.add_run(contact_line).font.size = Pt(9)

        doc.add_paragraph()  # boşluk

    def _add_summary(self, doc: Document, cv: TailoredCV) -> None:
        self._heading(doc, "PROFESSIONAL SUMMARY")
        p = doc.add_paragraph(cv.summary)
        p.runs[0].font.size = Pt(10)

    def _add_experience(self, doc: Document, cv: TailoredCV) -> None:
        self._heading(doc, "EXPERIENCE")
        for exp in CV_DATA["experience"]:
            company = exp["company"]
            bullets = cv.highlighted_bullets.get(company, exp["bullets"][:3])

            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_before = Pt(6)
            r = title_p.add_run(exp["title"])
            r.bold = True
            r.font.size = Pt(11)

            sub_p = doc.add_paragraph()
            sub_r = sub_p.add_run(f"{company}  |  {exp['period']}")
            sub_r.font.size = Pt(9)
            sub_r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
            sub_r.italic = True

            for bullet in bullets:
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(bullet).font.size = Pt(10)

    def _add_skills(self, doc: Document) -> None:
        self._heading(doc, "SKILLS")
        skills = CV_DATA["skills"]
        all_skills = (
            skills["leadership"][:4]
            + skills["product"][:4]
            + skills["crm_data"][:4]
            + skills["digital_transformation"][:3]
        )
        p = doc.add_paragraph()
        p.add_run("  ·  ".join(all_skills)).font.size = Pt(10)

    def _add_certifications(self, doc: Document) -> None:
        self._heading(doc, "CERTIFICATIONS")
        for cert in CV_DATA["certifications"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(cert).font.size = Pt(10)

    def _add_education(self, doc: Document) -> None:
        self._heading(doc, "EDUCATION")
        for edu in CV_DATA["education"]:
            p = doc.add_paragraph()
            r = p.add_run(f"{edu['degree']}  —  {edu['institution']}  |  {edu['period']}")
            r.font.size = Pt(10)
            if edu.get("note"):
                note_p = doc.add_paragraph()
                note_r = note_p.add_run(edu["note"])
                note_r.font.size = Pt(9)
                note_r.italic = True

    @staticmethod
    def _heading(doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x3b, 0x82, 0xf6)
        # Alt çizgi
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run("─" * 80)
        r2.font.size = Pt(6)
        r2.font.color.rgb = RGBColor(0xe2, 0xe8, 0xf0)
